"""Stage 4: Parse platform-exported .docx transcripts into TranscriptSegments.

Two timed formats are recognised — Microsoft Teams and Google Meet — plus an
untimed plain-paragraph fallback for hand-written Word transcripts. A document
that carries timecode structure we cannot align is **refused**, never degraded
to the untimed fallback: see ``DocxParseRefusedError``.
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from bristlenose.models import InputFile, TranscriptSegment
from bristlenose.utils.timecodes import parse_timecode

logger = logging.getLogger(__name__)

# Teams transcript patterns:
# "Speaker Name  00:01:23" or "Speaker Name 0:01:23"
_TEAMS_SPEAKER_LINE = re.compile(
    r"^(.+?)\s+(\d{1,2}:\d{2}(?::\d{2})?)\s*$"
)

# The same header, but insisting on the RUN OF SPACES Teams puts between the
# name and the timecode ("Martin Storey   0:04" — three, in the specimen).
#
# This exists only to license a ONE-header transcript. Two headers is weak
# evidence on its own because `_TEAMS_SPEAKER_LINE` is loose enough to match an
# ordinary sentence ending in a clock time, so the original guard demanded two
# of them. But the observed Teams export is a single 39-second turn, and one
# header is all a short interview ever has — the guard refused a valid file.
# The wide separator is the discriminator that a prose line ("Total runtime
# 1:45", one space) does not have.
_TEAMS_SPEAKER_LINE_WIDE = re.compile(
    r"^(.+?)\s{2,}(\d{1,2}:\d{2}(?::\d{2})?)\s*$"
)

# Teams with arrow notation: "Speaker Name  00:01:23 --> 00:01:45"
_TEAMS_SPEAKER_ARROW = re.compile(
    r"^(.+?)\s+(\d{1,2}:\d{2}:\d{2}[.,]\d{3})\s*-->\s*(\d{1,2}:\d{2}:\d{2}[.,]\d{3})\s*$"
)

# Just a timestamp line (no speaker)
_TIMESTAMP_ONLY = re.compile(
    r"^\s*(\d{1,2}:\d{2}(?::\d{2})?(?:[.,]\d{1,3})?)\s*$"
)

# ── Google Meet ────────────────────────────────────────────────────────────
#
# Meet saves a transcript to Drive as a Google Doc with TWO tabs — "Notes"
# (Gemini's AI summary of the meeting) and "Transcript" (the speech). Exporting
# that Doc to .docx FLATTENS BOTH TABS INTO ONE PARAGRAPH STREAM: python-docx
# sees the summary, the attendee list, and Gemini's own UI chrome as ordinary
# paragraphs, indistinguishable from speech by position alone.
#
# The only textual separators in the export are the tab headings "📝 Notes" and
# "📖 Transcript" — an emoji plus an English word that a non-English tenant
# localises. Matching those literals is the same fragility that made
# `_TEAMS_SUFFIX_RE` (s01_ingest) and Swift's `TeamsRecordingName` match only
# the fixtures invented alongside them.
#
# So we do NOT look for the tab boundary at all. We recognise SPEECH, by its
# structure: a paragraph that is nothing but an HH:MM:SS timecode, IMMEDIATELY
# followed by a paragraph of the form "Display Name: words". Only that pair
# yields a segment; every other paragraph is dropped. Gemini's prose carries no
# such pairing, so the Notes tab cannot be ingested as testimony — which is the
# failure that matters. The rule is also locale-independent: no emoji, no
# English word, nothing Google can rename.
#
# Grounded in a real specimen (Business Standard tenant, en-GB/BST, captured
# 15 Aug 2026) — tests/fixtures/platform-transcripts/gmeet-notes-by-gemini-2026-08-15.json.
_GMEET_TIMECODE_ONLY = re.compile(r"^(\d{1,2}:\d{2}:\d{2})$")

# "Martin Storey: Okay. Talking, talking, talking."
# The name may not contain a colon (Google display names don't), is length-
# bounded, and must be followed by non-empty speech. The mandatory preceding
# timecode does the real discriminating work; this is a sanity bound.
_GMEET_SPEAKER_TEXT = re.compile(r"^(?P<name>[^:]{1,80}):\s+(?P<text>\S.*)$")


class DocxParseRefusedError(RuntimeError):
    """A .docx carries timecode structure that could not be aligned to speech.

    Raised instead of degrading to :func:`_parse_plain_paragraphs`. A document
    with timecodes in it is a *timed transcript we failed to read*, not an
    untimed one: flattening it to paragraphs at ``start_time=0.0`` produces a
    report that renders cleanly while every timecode link, the coverage
    calculation and the player alignment are meaningless — and, for a Google
    Meet export, while AI summary prose sits in the corpus as quotable
    participant speech.

    The message is bristlenose-authored and names only the file basename — it
    never quotes document content, because it reaches ``Cause.message`` in
    ``pipeline-events.jsonl`` (a named re-identification surface).
    """


def parse_docx_file(input_file: InputFile) -> list[TranscriptSegment]:
    """Parse a .docx file (a Teams or Google Meet transcript export) into segments.

    Raises:
        DocxParseRefusedError: the document has timecodes but no readable speaker turns.
    """
    return _parse_docx(input_file.path)


def _parse_docx(path: Path) -> list[TranscriptSegment]:
    """Parse a platform-exported .docx file."""
    from docx import Document

    doc = Document(str(path))
    lines = _document_lines(doc)

    if not lines:
        logger.warning("Empty document: %s", path.name)
        return []

    # Google Meet first: its timecode+"Name: text" pairing is the more specific
    # shape, and a Meet document can incidentally satisfy the looser Teams
    # header pattern (the localised "Transcription ended after 00:00:39" trailer
    # matches `_TEAMS_SPEAKER_LINE`, for one).
    meet_segments, meet_dropped = _try_parse_gmeet_format(lines)
    if meet_segments:
        logger.info(
            "Parsed %d segments from Google Meet format (%d non-speech "
            "lines dropped, incl. any Gemini summary): %s",
            len(meet_segments),
            meet_dropped,
            path.name,
        )
        return _merge_adjacent_segments(meet_segments)

    segments = _try_parse_teams_format(lines)
    if segments:
        logger.info("Parsed %d segments from Teams format: %s", len(segments), path.name)
        return _merge_adjacent_segments(segments)

    # Timecodes present but unalignable is a parse REFUSAL, not plain text.
    # A line that is nothing but a timecode is strong evidence of a timed
    # transcript; incidental clock times inside prose ("we met at 14:30") are
    # not, because they never occupy a line alone.
    if any(_TIMESTAMP_ONLY.match(line) for line in lines):
        raise DocxParseRefusedError(
            f"{path.name} contains transcript timecodes but no readable speaker "
            f"turns. Export the transcript as .vtt, .srt or .txt instead."
        )

    # No timecode evidence at all — e.g. a transcription agency's Word file, or
    # a hand-typed one. Untimed paragraphs are the honest reading of that.
    #
    # KNOWN GAP: a Google Meet Doc whose Transcript tab is empty is also
    # untimed, and lands here — so its Gemini summary would be ingested as
    # speech. Closing that needs either the localised "📝 Notes" heading (the
    # fragility this parser deliberately avoids) or refusing every untimed
    # .docx, which would also reject legitimate agency transcripts. Left open
    # pending a product decision; the timed case above is the one observed in
    # the field.
    logger.info("Parsing as plain text (no timecodes): %s", path.name)
    return _parse_plain_paragraphs(lines)


def _document_lines(doc: object) -> list[str]:
    """Flatten the document to one entry per LINE, not per paragraph.

    A paragraph is a container, not a unit of transcript structure, and the two
    platforms disagree about how much they put in one. Meet writes one turn per
    paragraph; **Teams packs an entire turn — the "Name  0:04" header and every
    line of speech under it — into a single paragraph separated by ``<w:br/>``**,
    which python-docx returns as ``\\n`` inside ``paragraph.text``.

    So a parser that treats a paragraph as a line sees the real Teams export as
    one unmatched blob per turn, recognises nothing, and falls through to the
    untimed path — every segment at 0.0 with no speaker, which is the same
    silent mis-parse this module exists to prevent. Measured on the observed
    specimen (business tenant, 15 Aug 2026): the utterance paragraph carries 8
    line breaks, and the pre-split parser produced 5 untimed, unattributed
    segments from it.

    Splitting is safe for the other paths rather than merely tolerable: a Meet
    transcript turn carries no line break (measured — a 240-character
    multi-sentence turn had none), so Meet is unchanged in the transcript
    region, and a bare-timecode line that shared a paragraph with its speech
    would now be seen rather than missed.
    """
    lines: list[str] = []
    for para in doc.paragraphs:  # type: ignore[attr-defined]
        for raw in para.text.split("\n"):
            text = raw.strip()
            if text:
                lines.append(text)
    return lines


def _try_parse_gmeet_format(
    lines: list[str],
) -> tuple[list[TranscriptSegment], int]:
    """Try to parse lines as a Google Meet transcript export.

    Recognises speech structurally — an ``HH:MM:SS``-only line immediately
    followed by ``Name: text`` — and drops everything else. See the module-level
    Google Meet note for why this is boundary-free rather than heading-matched.

    Returns:
        ``(segments, dropped_line_count)``. Empty segments means "not this
        format"; the caller falls through.
    """
    segments: list[TranscriptSegment] = []
    consumed = 0
    i = 0

    while i < len(lines):
        ts_match = _GMEET_TIMECODE_ONLY.match(lines[i])
        if ts_match and i + 1 < len(lines):
            speech = _GMEET_SPEAKER_TEXT.match(lines[i + 1])
            if speech:
                start = parse_timecode(ts_match.group(1))
                segments.append(
                    TranscriptSegment(
                        start_time=start,
                        # Meet gives no turn end. Mirroring the Teams simple
                        # format (end == start) keeps `_merge_adjacent_segments`
                        # behaving identically across both platforms.
                        end_time=start,
                        text=speech.group("text").strip(),
                        speaker_label=speech.group("name").strip(),
                        source="docx",
                    )
                )
                consumed += 2
                i += 2
                continue
        i += 1

    return segments, len(lines) - consumed


def _try_parse_teams_format(lines: list[str]) -> list[TranscriptSegment] | None:
    """Try to parse lines as a Teams transcript.

    Teams transcripts alternate between speaker+timestamp lines and text lines.
    Both live inside ONE paragraph in a real export, so this must be given the
    line-split stream from `_document_lines` — handing it `doc.paragraphs`
    recognises nothing. Returns None if the format doesn't match.
    """
    segments: list[TranscriptSegment] = []
    current_speaker: str | None = None
    current_start: float | None = None
    current_end: float | None = None
    current_texts: list[str] = []
    matched_headers = 0
    strong_headers = 0  # arrow notation, or a name/timecode gap of 2+ spaces

    for para in lines:
        # Try matching speaker + timestamp (arrow format)
        arrow_match = _TEAMS_SPEAKER_ARROW.match(para)
        if arrow_match:
            # Flush previous segment
            if current_speaker is not None and current_texts:
                segments.append(_build_segment(
                    current_speaker, current_start, current_end, current_texts
                ))

            current_speaker = arrow_match.group(1).strip()
            current_start = parse_timecode(arrow_match.group(2))
            current_end = parse_timecode(arrow_match.group(3))
            current_texts = []
            matched_headers += 1
            strong_headers += 1  # "-->" is unambiguous; one is enough
            continue

        # Try matching speaker + timestamp (simple format)
        speaker_match = _TEAMS_SPEAKER_LINE.match(para)
        if speaker_match:
            # Flush previous segment
            if current_speaker is not None and current_texts:
                segments.append(_build_segment(
                    current_speaker, current_start, current_end, current_texts
                ))

            current_speaker = speaker_match.group(1).strip()
            current_start = parse_timecode(speaker_match.group(2))
            current_end = None
            if _TEAMS_SPEAKER_LINE_WIDE.match(para):
                strong_headers += 1
            current_texts = []
            matched_headers += 1
            continue

        # Try timestamp-only line
        ts_match = _TIMESTAMP_ONLY.match(para)
        if ts_match:
            if current_speaker is not None and current_texts:
                segments.append(_build_segment(
                    current_speaker, current_start, current_end, current_texts
                ))
                current_texts = []
            current_start = parse_timecode(ts_match.group(1))
            current_end = None
            continue

        # Otherwise it's a text line — add to current segment
        if current_speaker is not None:
            current_texts.append(para)

    # Flush final segment
    if current_speaker is not None and current_texts:
        segments.append(_build_segment(
            current_speaker, current_start, current_end, current_texts
        ))

    # Two loose headers, or one unambiguous one. A header with no speech under
    # it produced no segment, so `segments` non-empty is part of the evidence.
    if not segments:
        return None
    if matched_headers < 2 and strong_headers < 1:
        return None

    return segments


def _build_segment(
    speaker: str,
    start: float | None,
    end: float | None,
    texts: list[str],
) -> TranscriptSegment:
    """Build a TranscriptSegment from accumulated data."""
    text = " ".join(texts)
    return TranscriptSegment(
        start_time=start or 0.0,
        end_time=end or (start or 0.0),
        text=text,
        speaker_label=speaker,
        source="docx",
    )


def _parse_plain_paragraphs(paragraphs: list[str]) -> list[TranscriptSegment]:
    """Parse paragraphs with no timecode information.

    Each paragraph becomes a segment at time 0.0.
    """
    segments: list[TranscriptSegment] = []
    for i, para in enumerate(paragraphs):
        segments.append(
            TranscriptSegment(
                start_time=0.0,
                end_time=0.0,
                text=para,
                source="docx",
            )
        )
    return segments


def _merge_adjacent_segments(
    segments: list[TranscriptSegment],
    max_gap: float = 5.0,
) -> list[TranscriptSegment]:
    """Merge consecutive segments from the same speaker within max_gap seconds."""
    if not segments:
        return []

    merged: list[TranscriptSegment] = [segments[0].model_copy()]

    for seg in segments[1:]:
        prev = merged[-1]
        same_speaker = (
            prev.speaker_label is not None
            and prev.speaker_label == seg.speaker_label
        )
        close_enough = (seg.start_time - prev.end_time) <= max_gap

        if same_speaker and close_enough:
            prev.end_time = max(prev.end_time, seg.end_time)
            prev.text = f"{prev.text} {seg.text}"
        else:
            merged.append(seg.model_copy())

    return merged
