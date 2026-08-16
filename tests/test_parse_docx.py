"""Stage 4 .docx transcript parsing.

The Google Meet cases are pinned to an OBSERVED specimen — see
``tests/fixtures/platform-transcripts/gmeet-notes-by-gemini-2026-08-15.json``
and the "real specimens" note in ``tests/test_ingest.py``. Google Meet's .docx
export flattens the Gemini "Notes" tab and the "Transcript" tab into one
paragraph stream, so the interesting assertions are about what must NOT be
ingested, not only about what must.
"""

from __future__ import annotations

import asyncio
import json
from datetime import datetime
from pathlib import Path

import pytest

from bristlenose.models import FileType, InputFile, InputSession
from bristlenose.stages.s04_parse_docx import (
    DocxParseRefusedError,
    parse_docx_file,
)

FIXTURE = (
    Path(__file__).parent
    / "fixtures"
    / "platform-transcripts"
    / "gmeet-notes-by-gemini-2026-08-15.json"
)


def _specimen() -> dict:
    return json.loads(FIXTURE.read_text(encoding="utf-8"))


def _make_docx(tmp_path: Path, paragraphs: list[str], name: str = "t.docx") -> InputFile:
    """Write *paragraphs* as a .docx and wrap it as an InputFile.

    The specimen is the observed *paragraph stream* — that is what python-docx
    yields from Google's export and what the parser consumes. Rebuilding the
    container here keeps the real ``Document(...)`` read path under test without
    committing a binary that only resembles Google's.
    """
    from docx import Document

    doc = Document()
    for para in paragraphs:
        doc.add_paragraph(para)
    path = tmp_path / name
    doc.save(str(path))
    stat = path.stat()
    return InputFile(
        path=path,
        file_type=FileType.DOCX,
        created_at=datetime.fromtimestamp(stat.st_mtime),
        size_bytes=stat.st_size,
    )


# ── Google Meet: the observed specimen ─────────────────────────────────────


class TestGoogleMeetSpecimen:
    """Pinned to a real Business Standard tenant capture, 15 Aug 2026."""

    def test_only_the_speech_paragraph_becomes_a_segment(self, tmp_path: Path) -> None:
        """19 paragraphs in, exactly 1 segment out — the rest is not speech."""
        spec = _specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))

        assert len(segments) == spec["expected"]["segment_count"]
        assert segments[0].speaker_label == spec["expected"]["speaker_label"]
        assert segments[0].start_time == spec["expected"]["start_time"]
        assert segments[0].text.startswith("Okay. Talking, talking, talking.")

    def test_speaker_name_is_attribution_not_text(self, tmp_path: Path) -> None:
        """"Martin Storey:" must become the label, not stay embedded in speech."""
        spec = _specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))

        assert "Martin Storey:" not in segments[0].text

    def test_gemini_summary_never_enters_the_corpus(self, tmp_path: Path) -> None:
        """The AI summary tab must not become quotable participant testimony.

        This is the consequence that matters most: an LLM paraphrase of what
        people said, extracted downstream as a verbatim quote and presented in
        the report as testimony.
        """
        spec = _specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))
        corpus = " ".join(s.text for s in segments)

        assert "summary wasn't produced" not in corpus
        assert "Details weren't produced" not in corpus
        # Gemini's own UI chrome is quotable too if it lands in the corpus.
        assert "Gemini" not in corpus
        assert "short survey" not in corpus

    def test_attendee_email_never_enters_the_corpus(self, tmp_path: Path) -> None:
        """Stage 7 PII redaction is off by default — this must not reach it."""
        spec = _specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))
        corpus = " ".join(s.text for s in segments)

        assert "@" not in corpus

    def test_trailer_and_disclaimer_are_dropped(self, tmp_path: Path) -> None:
        """Google's own boilerplate is not something a participant said."""
        spec = _specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))
        corpus = " ".join(s.text for s in segments)

        assert "Transcription ended after" not in corpus
        assert "computer generated" not in corpus

    def test_timecode_is_real_not_zero(self, tmp_path: Path) -> None:
        """The pre-fix parser put every paragraph at 0.0, which silently broke
        timecode links, transcript coverage and player alignment."""
        spec = _specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))

        assert segments[0].start_time == 11.0


class TestGoogleMeetStructure:
    """The pairing rule, exercised beyond the single-turn specimen."""

    def test_multiple_turns(self, tmp_path: Path) -> None:
        paragraphs = [
            "📖 Transcript",
            "00:00:11",
            "Martin Storey: First thing I noticed was the price.",
            "00:00:25",
            "Ana Ferrer: That surprised me too.",
            "00:01:40",
            "Martin Storey: And the checkout took forever.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert [s.speaker_label for s in segments] == [
            "Martin Storey",
            "Ana Ferrer",
            "Martin Storey",
        ]
        assert [s.start_time for s in segments] == [11.0, 25.0, 100.0]

    def test_speech_containing_a_colon_keeps_its_colon(self, tmp_path: Path) -> None:
        """Only the first colon delimits the name."""
        paragraphs = [
            "00:00:11",
            "Martin Storey: I told them: it has to ship Friday.",
            "00:00:30",
            "Ana Ferrer: Agreed.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert segments[0].speaker_label == "Martin Storey"
        assert segments[0].text == "I told them: it has to ship Friday."

    def test_colon_prose_without_a_preceding_timecode_is_not_speech(
        self, tmp_path: Path
    ) -> None:
        """The timecode pairing is what stops Gemini prose being read as a turn.

        "Key decision: ship on Friday" is shaped exactly like a speaker line.
        Only its lack of a preceding timecode keeps it out of the corpus.
        """
        paragraphs = [
            "📝 Notes",
            "Key decision: ship on Friday",
            "Next step: Ana Ferrer to confirm pricing",
            "📖 Transcript",
            "00:00:11",
            "Martin Storey: The price was the thing.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert len(segments) == 1
        assert "Key decision" not in segments[0].text
        assert "Next step" not in segments[0].text

    def test_timecode_without_a_speaker_line_is_refused(self, tmp_path: Path) -> None:
        """Timecodes we cannot align to a speaker are a refusal, not plain text."""
        paragraphs = [
            "📖 Transcript",
            "00:00:11",
            "Some line with no speaker prefix at all",
            "00:00:30",
            "Another one",
        ]
        with pytest.raises(DocxParseRefusedError):
            parse_docx_file(_make_docx(tmp_path, paragraphs))


# ── Teams: unchanged behaviour ─────────────────────────────────────────────


class TestTeamsFormat:
    """The pre-existing Teams path must be unaffected by the Meet branch."""

    def test_speaker_plus_timecode_headers(self, tmp_path: Path) -> None:
        paragraphs = [
            "Martin Storey   0:00:11",
            "First thing I noticed was the price.",
            "Ana Ferrer   0:00:25",
            "That surprised me too.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert [s.speaker_label for s in segments] == ["Martin Storey", "Ana Ferrer"]
        assert [s.start_time for s in segments] == [11.0, 25.0]
        assert segments[0].text == "First thing I noticed was the price."

    def test_arrow_notation(self, tmp_path: Path) -> None:
        paragraphs = [
            "Martin Storey   00:00:11.000 --> 00:00:20.000",
            "First thing I noticed was the price.",
            "Ana Ferrer   00:00:25.000 --> 00:00:31.000",
            "That surprised me too.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert [s.speaker_label for s in segments] == ["Martin Storey", "Ana Ferrer"]
        assert segments[0].end_time == 20.0


# ── The fallback boundary ──────────────────────────────────────────────────


class TestPlainAndRefusal:
    def test_untimed_word_document_still_parses_as_paragraphs(
        self, tmp_path: Path
    ) -> None:
        """A hand-typed Word transcript has no timecodes — untimed is honest."""
        paragraphs = [
            "Interview with participant one",
            "They said the price was the first thing they noticed.",
            "Then they mentioned checkout was slow.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert len(segments) == 3
        assert all(s.start_time == 0.0 for s in segments)

    def test_incidental_clock_time_in_prose_does_not_trigger_refusal(
        self, tmp_path: Path
    ) -> None:
        """Only a paragraph that is *nothing but* a timecode is structural evidence."""
        paragraphs = [
            "Interview with participant one",
            "We met at 14:30 and talked for an hour.",
            "They said the price was the first thing they noticed.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert len(segments) == 3

    def test_empty_document_returns_nothing(self, tmp_path: Path) -> None:
        assert parse_docx_file(_make_docx(tmp_path, [])) == []


# ── The refusal must reach the run, not vanish ─────────────────────────────


class TestRefusalIsStatedNotDropped:
    """A refusal that nobody records is a worse bug than the mis-parse.

    ``docs/design-cloud-import.md``: "a parse refusal must produce a stated row,
    never a silent drop." A docx-bearing session has ``has_existing_transcript``
    set, so s02 skips audio extraction and leaves ``audio_path`` None — which
    also excludes it from ``needs_transcription``. Without a recorded failure it
    would leave no trace anywhere in the run.
    """

    def _session(self, tmp_path: Path, paragraphs: list[str]) -> InputSession:
        return InputSession(
            session_id="s1",
            session_number=1,
            participant_id="p1",
            participant_number=1,
            files=[_make_docx(tmp_path, paragraphs, "meet-transcript.docx")],
            has_existing_transcript=True,
            session_date=datetime(2026, 8, 15),
        )

    def _gather(self, session: InputSession):
        from bristlenose.config import BristlenoseSettings
        from bristlenose.pipeline import Pipeline

        settings = BristlenoseSettings(skip_transcription=True)
        pipeline = Pipeline(settings)
        return asyncio.run(pipeline._gather_all_segments([session]))

    def test_refusal_records_a_stage_failure(self, tmp_path: Path) -> None:
        unalignable = ["📖 Transcript", "00:00:11", "no speaker prefix here"]
        segments, outcome = self._gather(self._session(tmp_path, unalignable))

        assert segments == {}
        assert outcome.attempted == 1, "a failed attempt is still an attempt"
        assert outcome.succeeded == 0
        assert len(outcome.failed) == 1

        failure = outcome.failed[0]
        assert failure.session_id == "s1"
        assert failure.source_file == "meet-transcript.docx"
        assert failure.cause.stage == "s04_parse_docx"

    def test_failure_message_quotes_no_document_content(self, tmp_path: Path) -> None:
        """``Cause.message`` lands in pipeline-events.jsonl — a re-id surface."""
        unalignable = [
            "📖 Transcript",
            "00:00:11",
            "Bruce mentioned his diagnosis at the Watford clinic",
        ]
        _, outcome = self._gather(self._session(tmp_path, unalignable))

        message = outcome.failed[0].cause.message or ""
        assert "meet-transcript.docx" in message
        assert "Watford" not in message
        assert "diagnosis" not in message


# ── Teams: the observed specimen ───────────────────────────────────────────
#
# Every Teams case ABOVE this line is CONSTRUCTED — one line per paragraph,
# which is not what Teams writes. They passed against a parser that could not
# read a single real Teams export, which is the exact shape of the two
# six-month bugs recorded in docs/design-cloud-import.md §6.
#
# Cases below are pinned to a file observed on a real tenant.

TEAMS_FIXTURE = (
    Path(__file__).parent
    / "fixtures"
    / "platform-transcripts"
    / "teams-meeting-recording-2026-08-15.json"
)


def _teams_specimen() -> dict:
    return json.loads(TEAMS_FIXTURE.read_text(encoding="utf-8"))


class TestTeamsObservedSpecimen:
    """Business tenant, captured 15 Aug 2026 — a single 39-second turn."""

    def test_the_turn_is_timed_and_attributed(self, tmp_path: Path) -> None:
        """Pre-fix this file produced 5 segments, all 0.0, all unattributed."""
        spec = _teams_specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))

        assert len(segments) == spec["expected"]["segment_count"]
        assert segments[0].speaker_label == spec["expected"]["speaker_label"]
        assert segments[0].start_time == spec["expected"]["start_time"]
        assert segments[0].text.startswith(spec["expected"]["text_starts_with"])

    def test_the_specimen_really_does_pack_a_turn_into_one_paragraph(self) -> None:
        """Guard the guard: if this stops holding, the test above proves nothing.

        The whole Teams bug is that the speaker header and its speech share a
        paragraph. A fixture edited into one-line-per-paragraph would still pass
        the assertions above while testing nothing that failed in the field.
        """
        spec = _teams_specimen()
        multiline = [p for p in spec["paragraphs"] if "\n" in p]

        assert multiline, "specimen no longer carries embedded line breaks"
        assert any(p.count("\n") >= 7 for p in multiline)

    def test_session_furniture_is_not_speech(self, tmp_path: Path) -> None:
        """The export's own header block is not something a participant said."""
        spec = _teams_specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))
        corpus = " ".join(s.text for s in segments)

        assert "Meeting Recording" not in corpus  # the filename banner
        assert "06:07pm" not in corpus  # the date line
        assert "started transcription" not in corpus

    def test_stop_marker_is_a_known_residue(self, tmp_path: Path) -> None:
        """KNOWN GAP, pinned so it is visible rather than silent.

        Teams closes a transcript with "<Name> stopped transcription", which
        trails the final turn and is currently absorbed into it. Stripping it
        needs either the English literal — the localisation fragility this
        module deliberately avoids — or a rule that risks eating a genuine last
        line beginning with the speaker's name. Deferred pending a multi-turn
        specimen. If a future change fixes it, invert this assertion.
        """
        spec = _teams_specimen()
        segments = parse_docx_file(_make_docx(tmp_path, spec["paragraphs"]))

        assert "stopped transcription" in segments[-1].text


class TestParagraphIsNotTheUnitOfStructure:
    """The two platforms disagree about how much goes in one paragraph."""

    def test_teams_turn_split_across_line_breaks(self, tmp_path: Path) -> None:
        """One paragraph, header and speech separated by <w:br/>."""
        paragraphs = [
            "Martin Storey   0:04\nThe price was the first thing.\nThen checkout.",
            "Ana Ferrer   0:31\nThat surprised me too.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert [s.speaker_label for s in segments] == ["Martin Storey", "Ana Ferrer"]
        assert [s.start_time for s in segments] == [4.0, 31.0]
        assert segments[0].text == "The price was the first thing. Then checkout."

    def test_single_wide_header_is_enough(self, tmp_path: Path) -> None:
        """A short interview has one turn; two headers is not always available."""
        paragraphs = ["Martin Storey   0:04\nOnly one turn in this recording."]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert len(segments) == 1
        assert segments[0].speaker_label == "Martin Storey"
        assert segments[0].start_time == 4.0

    def test_prose_ending_in_a_clock_time_is_not_a_speaker(
        self, tmp_path: Path
    ) -> None:
        """The single-header licence rests on Teams' run of spaces.

        "Total runtime 1:45" is one space and must not become a speaker, or an
        untimed agency transcript silently acquires a fabricated participant.
        """
        paragraphs = [
            "Interview with participant one",
            "Total runtime 1:45",
            "They said the price was the first thing they noticed.",
        ]
        segments = parse_docx_file(_make_docx(tmp_path, paragraphs))

        assert all(s.speaker_label is None for s in segments)
        assert len(segments) == 3
